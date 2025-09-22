import customtkinter
from tkinter import filedialog, messagebox
import pandas as pd
import openpyxl
import docx
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import shutil
import threading
import os
import math

# ==============================================================================
# --- FUNGSI BANTU (Berlaku untuk kedua tugas) ---
# ==============================================================================
def sanitize_sheet_name(name):
    name = str(name).strip(); return re.sub(r'[\\/*?:"<>|]', '', name)[:31] if name else "Data_Kosong"
def set_cell_border(cell, **kwargs):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom', 'start', 'end'):
        if edge_data := kwargs.get(edge):
            tag = f'w:{edge}'; border = OxmlElement(tag)
            for k, v in edge_data.items(): border.set(qn(f'w:{k}'), str(v))
            tcBorders.append(border)
    tcPr.append(tcBorders)
def set_cell_margins(cell, **kwargs):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = OxmlElement('w:tcMar')
    for edge in ('top', 'left', 'bottom', 'right'):
        if margin_data := kwargs.get(edge):
            tag = f'w:{edge}'; margin = OxmlElement(tag)
            margin.set(qn('w:w'), str(margin_data.twips)); margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
    tcPr.append(tcMar)

# ==============================================================================
# --- KELAS APLIKASI UTAMA ---
# ==============================================================================
class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()
        self.title("Aplikasi Otomatisasi Dokumen v2.2 (Responsif)")
        self.geometry("800x800")
        
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(2, weight=1)

        # Variabel State
        self.task_choice = customtkinter.StringVar(value="Label")
        self.data_file_path = ""
        self.template_file_path = ""
        self.label_data_sheet = customtkinter.StringVar()
        self.label_column_map = {'customer': customtkinter.StringVar(),'alamat': customtkinter.StringVar(),'pic': customtkinter.StringVar(),'telp': customtkinter.StringVar()}
        self.receipt_data_sheet = customtkinter.StringVar()
        self.receipt_template_sheet = customtkinter.StringVar()
        self.receipt_column_map = {'plat': customtkinter.StringVar(), 'pic': customtkinter.StringVar(), 'customer': customtkinter.StringVar()}

        # WIDGET UTAMA
        self.title_label = customtkinter.CTkLabel(self, text="Aplikasi Otomatisasi Dokumen", font=customtkinter.CTkFont(size=20, weight="bold"))
        self.title_label.grid(row=0, column=0, padx=20, pady=20)
        self.task_frame = customtkinter.CTkFrame(self)
        self.task_frame.grid(row=1, column=0, padx=20, pady=10, sticky="ew")
        customtkinter.CTkLabel(self.task_frame, text="Pilih Tugas:").pack(side="left", padx=10, pady=10)
        self.radio_label = customtkinter.CTkRadioButton(self.task_frame, text="Buat Label (Word)", variable=self.task_choice, value="Label", command=self.toggle_task_view)
        self.radio_label.pack(side="left", padx=10, pady=10)
        self.radio_receipt = customtkinter.CTkRadioButton(self.task_frame, text="Isi Tanda Terima (Excel)", variable=self.task_choice, value="Receipt", command=self.toggle_task_view)
        self.radio_receipt.pack(side="left", padx=10, pady=10)

        # Frame untuk setiap tugas
        self.label_frame = customtkinter.CTkFrame(self, fg_color="transparent")
        self.label_frame.grid_columnconfigure(1, weight=1)
        self.receipt_frame = customtkinter.CTkFrame(self, fg_color="transparent")
        self.receipt_frame.grid_columnconfigure(1, weight=1)

        self.build_label_ui()
        self.build_receipt_ui()

        self.generate_button = customtkinter.CTkButton(self, text="GENERATE", command=self.start_generation, height=40, font=customtkinter.CTkFont(size=14, weight="bold"))
        self.generate_button.grid(row=3, column=0, padx=20, pady=10, sticky="ew")
        self.status_label = customtkinter.CTkLabel(self, text="Selamat datang! Silakan pilih tugas.", wraplength=750)
        self.status_label.grid(row=4, column=0, padx=20, pady=10)
        
        self.toggle_task_view()

    def toggle_task_view(self):
        if self.task_choice.get() == "Label":
            self.receipt_frame.grid_remove()
            self.label_frame.grid(row=2, column=0, padx=20, pady=10, sticky="nsew")
        else:
            self.label_frame.grid_remove()
            self.receipt_frame.grid(row=2, column=0, padx=20, pady=10, sticky="nsew")

    def start_generation(self):
        self.generate_button.configure(state="disabled")
        self.status_label.configure(text="Memproses... mohon tunggu...", text_color="yellow")
        
        task = self.task_choice.get()
        if task == "Label":
            if not self.data_file_path or any(not v.get() or v.get() == "-" for v in self.label_column_map.values()):
                messagebox.showerror("Input Tidak Lengkap", "Untuk 'Buat Label', harap pilih File Data dan cocokkan semua kolom.")
                self.generate_button.configure(state="normal"); self.status_label.configure(text="Proses dibatalkan.", text_color="gray")
                return
            output_path = filedialog.asksaveasfilename(title="Simpan File Label Word", defaultextension=".docx", filetypes=(("Word Document", "*.docx"),))
            if not output_path: self.generate_button.configure(state="normal"); self.status_label.configure(text="Proses dibatalkan.", text_color="gray"); return
            thread = threading.Thread(target=self.run_label_generation, args=(output_path,))
            thread.start()
        else:
            if not self.data_file_path or not self.template_file_path or any(not v.get() or v.get() == "-" for v in self.receipt_column_map.values()):
                messagebox.showerror("Input Tidak Lengkap", "Untuk 'Isi Tanda Terima', harap pilih File Data, File Template, dan cocokkan semua kolom.")
                self.generate_button.configure(state="normal"); self.status_label.configure(text="Proses dibatalkan.", text_color="gray")
                return
            output_path = filedialog.asksaveasfilename(title="Simpan File Tanda Terima Excel", defaultextension=".xlsx", filetypes=(("Excel Files", "*.xlsx"),))
            if not output_path: self.generate_button.configure(state="normal"); self.status_label.configure(text="Proses dibatalkan.", text_color="gray"); return
            thread = threading.Thread(target=self.run_receipt_generation, args=(output_path,))
            thread.start()

    def build_label_ui(self):
        customtkinter.CTkLabel(self.label_frame, text="Pengaturan untuk Membuat Label (Word)", font=customtkinter.CTkFont(weight="bold")).grid(row=0, column=0, columnspan=2, padx=10, pady=(10,0), sticky="w")
        customtkinter.CTkLabel(self.label_frame, text="1. Pilih File Data Master:").grid(row=1, column=0, columnspan=2, padx=10, pady=(10,0), sticky="w")
        self.label_file_button = customtkinter.CTkButton(self.label_frame, text="Pilih File...", command=lambda: self.select_file('data', self.label_file_label, self.label_sheet_menu, self.load_label_columns))
        self.label_file_button.grid(row=2, column=0, padx=10, pady=5, sticky="w")
        self.label_file_label = customtkinter.CTkLabel(self.label_frame, text="Belum ada file dipilih", text_color="gray", anchor="w")
        self.label_file_label.grid(row=2, column=1, padx=10, pady=5, sticky="ew")
        customtkinter.CTkLabel(self.label_frame, text="2. Pilih Sheet Data:").grid(row=3, column=0, columnspan=2, padx=10, pady=(10,0), sticky="w")
        self.label_sheet_menu = customtkinter.CTkOptionMenu(self.label_frame, variable=self.label_data_sheet, values=["-"], command=self.load_label_columns)
        self.label_sheet_menu.grid(row=4, column=0, columnspan=2, padx=10, pady=5, sticky="ew")
        customtkinter.CTkLabel(self.label_frame, text="3. Cocokkan Kolom:").grid(row=5, column=0, columnspan=2, padx=10, pady=(10,0), sticky="w")
        map_frame = customtkinter.CTkFrame(self.label_frame); map_frame.grid(row=6, column=0, columnspan=2, sticky="ew", padx=10, pady=5)
        map_frame.grid_columnconfigure(1, weight=1)
        customtkinter.CTkLabel(map_frame, text="Customer Name:").grid(row=0, column=0, padx=10, pady=5, sticky="e")
        self.label_customer_menu = customtkinter.CTkOptionMenu(map_frame, variable=self.label_column_map['customer'], values=["-"])
        self.label_customer_menu.grid(row=0, column=1, padx=10, pady=5, sticky="ew")
        customtkinter.CTkLabel(map_frame, text="Alamat:").grid(row=1, column=0, padx=10, pady=5, sticky="e")
        self.label_alamat_menu = customtkinter.CTkOptionMenu(map_frame, variable=self.label_column_map['alamat'], values=["-"])
        self.label_alamat_menu.grid(row=1, column=1, padx=10, pady=5, sticky="ew")
        customtkinter.CTkLabel(map_frame, text="PIC:").grid(row=2, column=0, padx=10, pady=5, sticky="e")
        self.label_pic_menu = customtkinter.CTkOptionMenu(map_frame, variable=self.label_column_map['pic'], values=["-"])
        self.label_pic_menu.grid(row=2, column=1, padx=10, pady=5, sticky="ew")
        customtkinter.CTkLabel(map_frame, text="No. Telp:").grid(row=3, column=0, padx=10, pady=5, sticky="e")
        self.label_telp_menu = customtkinter.CTkOptionMenu(map_frame, variable=self.label_column_map['telp'], values=["-"])
        self.label_telp_menu.grid(row=3, column=1, padx=10, pady=5, sticky="ew")

    def load_label_columns(self, selected_sheet):
        menus = [self.label_customer_menu, self.label_alamat_menu, self.label_pic_menu, self.label_telp_menu]
        self.load_columns_generic(self.data_file_path, selected_sheet, menus, self.label_column_map)

    def run_label_generation(self, output_path):
        try:
            KOLOM_CUSTOMER=self.label_column_map['customer'].get(); KOLOM_ALAMAT=self.label_column_map['alamat'].get(); KOLOM_PIC=self.label_column_map['pic'].get(); KOLOM_TELP=self.label_column_map['telp'].get()
            df = pd.read_excel(self.data_file_path, sheet_name=self.label_data_sheet.get()).fillna('')
            df_unik = df.drop_duplicates(subset=[KOLOM_PIC, KOLOM_ALAMAT], keep='first').reset_index(drop=True)
            total_labels = len(df_unik); JUMLAH_KOLOM = 2; total_rows = math.ceil(total_labels / JUMLAH_KOLOM)
            doc = docx.Document(); section = doc.sections[0]; section.top_margin = Cm(1.1); section.bottom_margin = Cm(1.1); section.left_margin = Cm(1); section.right_margin = Cm(1)
            table = doc.add_table(rows=total_rows, cols=JUMLAH_KOLOM); table.autofit = False; table.allow_autofit = False
            tbl_pr = table._element.xpath('w:tblPr')[0]; tbl_cell_spacing = OxmlElement('w:tblCellSpacing'); tbl_cell_spacing.set(qn('w:w'), str(Cm(0.2).twips)); tbl_pr.append(tbl_cell_spacing)
            data_iterator = iter(df_unik.to_dict('records'))
            for row_idx in range(total_rows):
                tr = table.rows[row_idx]._tr; trPr = tr.get_or_add_trPr(); cantSplit = OxmlElement('w:cantSplit'); trPr.append(cantSplit)
                for col_idx in range(JUMLAH_KOLOM):
                    try:
                        data = next(data_iterator); cell = table.cell(row_idx, col_idx); cell.width = Cm(9.5)
                        border_style = {"sz": 6, "val": "single", "color": "#000000"}; set_cell_border(cell, top=border_style, bottom=border_style, start=border_style, end=border_style)
                        set_cell_margins(cell, top=Cm(0), bottom=Cm(0.1), left=Cm(0.1), right=Cm(0.1))
                        p = cell.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 0.95
                        run_to = p.add_run('To :\n'); run_to.font.name = 'Calibri'; run_to.font.size = Pt(12); run_to.bold = True
                        run_cust = p.add_run(f"{data[KOLOM_CUSTOMER]}\n"); run_cust.font.name = 'Calibri'; run_cust.font.size = Pt(12); run_cust.bold = True
                        run_addr = p.add_run(f"{data[KOLOM_ALAMAT]}\n"); run_addr.font.name = 'Calibri'; run_addr.font.size = Pt(12); run_addr.bold = False
                        run_pic = p.add_run(f"Up. Bpk/Ibu {data[KOLOM_PIC]} (Hp. {data[KOLOM_TELP]})"); run_pic.font.name = 'Calibri'; run_pic.font.size = Pt(12); run_pic.bold = True
                    except StopIteration: break
            doc.save(output_path)
            self.status_label.configure(text=f"Berhasil! File Label disimpan di:\n{output_path}", text_color="light green")
        except Exception as e:
            self.status_label.configure(text=f"Error: {e}", text_color="red"); messagebox.showerror("Error", f"Terjadi kesalahan saat membuat Label:\n\n{e}")
        finally: self.generate_button.configure(state="normal")
    
    def build_receipt_ui(self):
        customtkinter.CTkLabel(self.receipt_frame, text="Pengaturan untuk Mengisi Tanda Terima (Excel)", font=customtkinter.CTkFont(weight="bold")).grid(row=0, column=0, columnspan=2, padx=10, pady=(10,0), sticky="w")
        customtkinter.CTkLabel(self.receipt_frame, text="1. Pilih File Data Master:").grid(row=1, column=0, columnspan=2, padx=10, pady=(10,0), sticky="w")
        self.receipt_data_button = customtkinter.CTkButton(self.receipt_frame, text="Pilih File...", command=lambda: self.select_file('data', self.receipt_data_label, self.receipt_data_sheet_menu, self.load_receipt_columns))
        self.receipt_data_button.grid(row=2, column=0, padx=10, pady=5, sticky="w")
        self.receipt_data_label = customtkinter.CTkLabel(self.receipt_frame, text="Belum ada file dipilih", text_color="gray", anchor="w")
        self.receipt_data_label.grid(row=2, column=1, padx=10, pady=5, sticky="ew")
        customtkinter.CTkLabel(self.receipt_frame, text="2. Pilih File Template:").grid(row=3, column=0, columnspan=2, padx=10, pady=(10,0), sticky="w")
        self.receipt_template_button = customtkinter.CTkButton(self.receipt_frame, text="Pilih File...", command=lambda: self.select_file('template', self.receipt_template_label, self.receipt_template_sheet_menu))
        self.receipt_template_button.grid(row=4, column=0, padx=10, pady=5, sticky="w")
        self.receipt_template_label = customtkinter.CTkLabel(self.receipt_frame, text="Belum ada file dipilih", text_color="gray", anchor="w")
        self.receipt_template_label.grid(row=4, column=1, padx=10, pady=5, sticky="ew")
        customtkinter.CTkLabel(self.receipt_frame, text="3. Pilih Nama Sheet:").grid(row=5, column=0, columnspan=2, padx=10, pady=(10,0), sticky="w")
        self.receipt_data_sheet_menu = customtkinter.CTkOptionMenu(self.receipt_frame, variable=self.receipt_data_sheet, values=["- Data Master -"], command=self.load_receipt_columns)
        self.receipt_data_sheet_menu.grid(row=6, column=0, padx=10, pady=5, sticky="ew")
        self.receipt_template_sheet_menu = customtkinter.CTkOptionMenu(self.receipt_frame, variable=self.receipt_template_sheet, values=["- Template -"])
        self.receipt_template_sheet_menu.grid(row=6, column=1, padx=10, pady=5, sticky="ew")
        customtkinter.CTkLabel(self.receipt_frame, text="4. Cocokkan Kolom:").grid(row=7, column=0, columnspan=2, padx=10, pady=(10,0), sticky="w")
        map_frame = customtkinter.CTkFrame(self.receipt_frame); map_frame.grid(row=8, column=0, columnspan=2, sticky="ew", padx=10, pady=5)
        map_frame.grid_columnconfigure(1, weight=1)
        customtkinter.CTkLabel(map_frame, text="Plat Fix:").grid(row=0, column=0, padx=10, pady=5, sticky="e")
        self.receipt_plat_menu = customtkinter.CTkOptionMenu(map_frame, variable=self.receipt_column_map['plat'], values=["-"])
        self.receipt_plat_menu.grid(row=0, column=1, padx=10, pady=5, sticky="ew")
        customtkinter.CTkLabel(map_frame, text="PIC:").grid(row=1, column=0, padx=10, pady=5, sticky="e")
        self.receipt_pic_menu = customtkinter.CTkOptionMenu(map_frame, variable=self.receipt_column_map['pic'], values=["-"])
        self.receipt_pic_menu.grid(row=1, column=1, padx=10, pady=5, sticky="ew")
        customtkinter.CTkLabel(map_frame, text="Customer Name:").grid(row=2, column=0, padx=10, pady=5, sticky="e")
        self.receipt_customer_menu = customtkinter.CTkOptionMenu(map_frame, variable=self.receipt_column_map['customer'], values=["-"])
        self.receipt_customer_menu.grid(row=2, column=1, padx=10, pady=5, sticky="ew")

    def load_receipt_columns(self, selected_sheet):
        menus = [self.receipt_plat_menu, self.receipt_pic_menu, self.receipt_customer_menu]
        self.load_columns_generic(self.data_file_path, selected_sheet, menus, self.receipt_column_map)
        
    def run_receipt_generation(self, output_path):
        try:
            KOLOM_PLAT=self.receipt_column_map['plat'].get(); KOLOM_PIC=self.receipt_column_map['pic'].get(); KOLOM_CUSTOMER=self.receipt_column_map['customer'].get()
            shutil.copy(self.template_file_path, output_path)
            df = pd.read_excel(self.data_file_path, sheet_name=self.receipt_data_sheet.get()).fillna('')
            grouped = df.groupby([KOLOM_PIC, KOLOM_CUSTOMER])
            workbook = openpyxl.load_workbook(output_path)
            template_sheet_name = self.receipt_template_sheet.get()
            template_sheet = workbook[template_sheet_name]
            for (pic_name, customer_name), group_data in grouped:
                sheet_title = f"{pic_name} - {customer_name}"; new_sheet_name = sanitize_sheet_name(sheet_title)
                sheet_suffix = 1; temp_name = new_sheet_name
                while temp_name in workbook.sheetnames: sheet_suffix += 1; temp_name = f"{new_sheet_name[:28]}_{sheet_suffix}"
                new_sheet_name = temp_name
                new_sheet = workbook.copy_worksheet(template_sheet); new_sheet.title = new_sheet_name
                new_sheet['F36'] = f"Bapak/Ibu {pic_name}"; new_sheet['P36'] = f"Bapak/Ibu {pic_name}"
                new_sheet['F37'] = customer_name; new_sheet['P37'] = customer_name
                plat_list = group_data[KOLOM_PLAT].tolist()
                start_row_plat = 16
                for i, plat in enumerate(plat_list):
                    new_sheet[f'D{start_row_plat + i}'] = f"{i + 1})"; new_sheet[f'N{start_row_plat + i}'] = f"{i + 1})"
                    new_sheet[f'E{start_row_plat + i}'] = plat; new_sheet[f'O{start_row_plat + i}'] = plat
            if template_sheet_name in workbook.sheetnames: del workbook[template_sheet_name]
            workbook.save(output_path)
            self.status_label.configure(text=f"Berhasil! File Tanda Terima disimpan di:\n{output_path}", text_color="light green")
        except Exception as e:
            self.status_label.configure(text=f"Error: {e}", text_color="red"); messagebox.showerror("Error", f"Terjadi kesalahan saat membuat Tanda Terima:\n\n{e}")
        finally: self.generate_button.configure(state="normal")
            
    def select_file(self, file_type, label_widget, sheet_menu, callback=None):
        path = filedialog.askopenfilename(title=f"Pilih file {'Data Master' if file_type == 'data' else 'Template'}", filetypes=(("Excel Files", "*.xlsx;*.xls"),))
        if not path: return
        
        if file_type == 'data': self.data_file_path = path
        elif file_type == 'template': self.template_file_path = path
        
        label_widget.configure(text=os.path.basename(path), text_color="white") 
        self.load_sheets(path, sheet_menu, callback)

    def load_sheets(self, path, menu, callback=None):
        try:
            xls = pd.ExcelFile(path); sheet_names = xls.sheet_names; menu.configure(values=sheet_names)
            if sheet_names:
                menu.set(sheet_names[0])
                if callback: callback(sheet_names[0])
        except Exception as e: messagebox.showerror("Error", f"Gagal membaca sheet dari file.\n\n{e}")

    def load_columns_generic(self, file_path, selected_sheet, menu_list, column_map_dict):
        if not file_path: return
        try:
            df_headers = pd.read_excel(file_path, sheet_name=selected_sheet, nrows=0)
            column_headers = ["-"] + df_headers.columns.tolist()
            for menu in menu_list:
                menu.configure(values=column_headers)
            for var in column_map_dict.values():
                var.set("-")
        except Exception as e: messagebox.showerror("Error", f"Gagal membaca kolom dari sheet '{selected_sheet}'.\n\n{e}")

if __name__ == "__main__":
    app = App()
    app.mainloop()